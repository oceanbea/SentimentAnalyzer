import csv
from docx import Document

def copy_csv_to_docx(csv_file, docx_file, chunk_size=1000):
    document = Document()
    
    with open(csv_file, 'r', newline='', encoding='utf-8') as file:
        reader = csv.DictReader(file)
        headers = reader.fieldnames
        writer = csv.DictWriter(doc_file, fieldnames=headers)
        writer.writeheader()
        
        for i, row in enumerate(reader):
            if i % chunk_size == 0:
                document.save(docx_file)
                document = Document()
                document.add_paragraph('\n')  # Add a blank line between chunks
            
            summary = row.get('Summary', '')
            rate = row.get('Rate', '')
            
            document.add_paragraph(f'Summary: {summary}\nRate: {rate}\n')
    
    document.save(docx_file)

# Replace 'products.csv' and 'output.docx' with the actual names of your files
csv_file = 'Dataset.csv'
docx_file = 'output.docx'

copy_csv_to_docx(csv_file, docx_file)
